"""Tests for cover letter generation and application notes fallbacks."""

import pytest
from docx import Document

import jobapply.content
from job_search_apply import (
    _basic_cover_letter,
    _basic_notes,
    _ensure_cover_letter_docx,
    _save_cover_letter_docx,
)


@pytest.fixture
def cover_letter_dir(tmp_path, monkeypatch):
    """Point COVER_LETTER_DIR at tmp_path so no test touches real data."""
    monkeypatch.setattr(jobapply.content, "COVER_LETTER_DIR", tmp_path)
    return tmp_path


class TestBasicCoverLetter:
    def test_without_template(self, profile, job):
        profile.cover_letter_template = None
        result = _basic_cover_letter(job, profile)
        assert "Senior DevOps Engineer" in result
        assert "TechCo" in result
        assert "Test User" in result
        assert "test@example.com" in result

    def test_with_template(self, profile, job):
        profile.cover_letter_template = (
            "{DATE}\n{COMPANY}\nRE: {JOB_TITLE}\n\n"
            "Dear {HIRING_MANAGER_NAME},\nI want this job.\n"
            "---\nTEMPLATE INSTRUCTIONS FOR AI:\nFill in the placeholders."
        )
        result = _basic_cover_letter(job, profile)
        assert "TechCo" in result
        assert "Senior DevOps Engineer" in result
        assert "{COMPANY}" not in result
        assert "{JOB_TITLE}" not in result
        # AI instructions should be stripped
        assert "TEMPLATE INSTRUCTIONS" not in result

    def test_years_in_generic_letter(self, profile, job):
        profile.cover_letter_template = None
        result = _basic_cover_letter(job, profile)
        assert "9" in result  # years of experience


class TestBasicNotes:
    def test_contains_job_info(self, job, compat):
        result = _basic_notes(job, compat)
        assert "Senior DevOps Engineer" in result
        assert "TechCo" in result
        assert "https://www.linkedin.com/jobs/view/123" in result

    def test_contains_score(self, job, compat):
        result = _basic_notes(job, compat)
        assert "0.85" in result

    def test_contains_skills(self, job, compat):
        result = _basic_notes(job, compat)
        assert "Kubernetes" in result

    def test_contains_description_snippet(self, job, compat):
        result = _basic_notes(job, compat)
        assert "Kubernetes" in result or "DevOps" in result

    def test_truncates_long_description(self, job, compat):
        job["description"] = "x " * 500  # 1000 chars
        result = _basic_notes(job, compat)
        # Description should be truncated with ellipsis
        assert len(result) < 2000


class TestSaveCoverLetterDocx:
    def test_writes_docx_file(self, cover_letter_dir):
        result = _save_cover_letter_docx("Dear team,\n\nI am applying.", "li_abc123")
        assert result == cover_letter_dir / "li_abc123.docx"
        assert result.exists()

    def test_text_round_trips(self, cover_letter_dir):
        text = "Dear team,\n\nI am applying for the role.\nThanks,\nTest User"
        path = _save_cover_letter_docx(text, "li_roundtrip")
        doc = Document(str(path))
        assert [p.text for p in doc.paragraphs] == text.split("\n")

    def test_uses_calibri_11pt(self, cover_letter_dir):
        # Pins the Comeet-era formatting: Calibri 11pt Normal style
        path = _save_cover_letter_docx("Hello", "li_style")
        doc = Document(str(path))
        style = doc.styles["Normal"]
        assert style.font.name == "Calibri"
        assert style.font.size.pt == 11


class TestEnsureCoverLetterDocx:
    def test_converts_legacy_txt_to_docx(self, cover_letter_dir):
        # The Comeet regression: .txt uploads are silently rejected by some ATS
        txt_path = cover_letter_dir / "li_legacy.txt"
        txt_path.write_text("Dear team,\nlegacy letter")
        result = _ensure_cover_letter_docx(str(txt_path))
        assert result == str(cover_letter_dir / "li_legacy.docx")
        doc = Document(result)
        assert [p.text for p in doc.paragraphs] == ["Dear team,", "legacy letter"]

    def test_docx_path_returned_unchanged(self, cover_letter_dir):
        path = _save_cover_letter_docx("Hello", "li_already")
        assert _ensure_cover_letter_docx(str(path)) == str(path)

    def test_missing_file_returned_unchanged(self, cover_letter_dir):
        missing = cover_letter_dir / "li_missing.txt"
        assert _ensure_cover_letter_docx(str(missing)) == str(missing)
        assert not (cover_letter_dir / "li_missing.docx").exists()

    def test_existing_docx_sibling_not_overwritten(self, cover_letter_dir):
        txt_path = cover_letter_dir / "li_dup.txt"
        txt_path.write_text("txt content")
        docx_path = cover_letter_dir / "li_dup.docx"
        docx_path.write_bytes(b"sentinel")
        result = _ensure_cover_letter_docx(str(txt_path))
        assert result == str(docx_path)
        assert docx_path.read_bytes() == b"sentinel"

    def test_txt_outside_cover_letter_dir_returns_real_path(self, tmp_path, monkeypatch):
        # Fixed 2026-07-08: the converted .docx is written to
        # COVER_LETTER_DIR and the function now returns that real path
        # (it used to return the .txt's nonexistent sibling, breaking the
        # subsequent upload).
        cl_dir = tmp_path / "cover-letters"
        cl_dir.mkdir()
        monkeypatch.setattr(jobapply.content, "COVER_LETTER_DIR", cl_dir)
        elsewhere = tmp_path / "elsewhere"
        elsewhere.mkdir()
        txt_path = elsewhere / "li_stray.txt"
        txt_path.write_text("stray letter")
        result = _ensure_cover_letter_docx(str(txt_path))
        assert result == str(cl_dir / "li_stray.docx")
        assert (cl_dir / "li_stray.docx").exists()
